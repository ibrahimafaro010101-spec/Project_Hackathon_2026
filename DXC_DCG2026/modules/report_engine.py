# ============================================================
# report_engine.py
# Moteur de g√©n√©ration de rapports utilisant OpenAI
# ============================================================

from datetime import datetime
import markdown
from typing import Dict, List, Optional, Any
from io import BytesIO
import re
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfgen import canvas

    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False


class ReportEngine:
    """
    G√©n√©rateur de rapports intelligents avec OpenAI
    """

    def __init__(self, ai_client=None):
        """
        Initialise avec un client OpenAI
        """
        self.ai_client = ai_client
        self.using_openai = ai_client is not None
        self.report_history = []

    # --------------------------------------------------------
    # M√âTHODES PUBLIQUES - G√âN√âRATION
    # --------------------------------------------------------

    def generate_report(
            self,
            title: str,
            audience: str,
            sections: list,
            data_summary: dict,
            analysis_summary: str = "",
            model_results: Optional[Dict] = None,
            insights: Optional[List] = None,
            custom_instructions: str = "",
            detail_level: int = 3,
            language: str = "fr",
            include_visualizations: bool = True
    ) -> str:
        """
        G√©n√®re un rapport complet avec OpenAI
        """
        try:
            # Pr√©paration du contexte d√©taill√©
            data_context = self._prepare_data_context(
                title=title,
                audience=audience,
                sections=sections,
                data_summary=data_summary,
                analysis_summary=analysis_summary,
                model_results=model_results,
                insights=insights,
                custom_instructions=custom_instructions,
                detail_level=detail_level,
                language=language
            )

            # Log de la g√©n√©ration
            generation_id = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            print(f"üîß G√©n√©ration du rapport: {generation_id}")

            # G√©n√©ration avec OpenAI si disponible
            if self.ai_client and hasattr(self.ai_client, 'generate_text'):
                print("üéØ Utilisation d'OpenAI pour la g√©n√©ration du rapport")

                try:
                    # Construction du prompt expert
                    prompt = self._build_expert_prompt(data_context)

                    # Param√®tres de g√©n√©ration adapt√©s au rapport
                    generation_params = {
                        "model": "gpt-4-turbo-preview",
                        "temperature": 0.7,
                        "max_tokens": 4000,
                        "top_p": 0.95,
                        "frequency_penalty": 0.2,
                        "presence_penalty": 0.1
                    }

                    # G√©n√©ration du rapport
                    report_text = self.ai_client.generate_text(
                        prompt=prompt,
                        **generation_params
                    )

                    # Post-traitement et validation
                    report_text = self._post_process_report(report_text, data_context)

                    # Sauvegarde dans l'historique
                    self.report_history.append({
                        "id": generation_id,
                        "timestamp": datetime.now().isoformat(),
                        "title": title,
                        "audience": audience,
                        "sections": sections,
                        "data_summary": data_summary,
                        "report_preview": report_text[:500] + "..." if len(report_text) > 500 else report_text
                    })

                    print(f"‚úÖ Rapport g√©n√©r√© avec succ√®s ({len(report_text)} caract√®res)")
                    return report_text

                except Exception as e:
                    print(f"‚ö†Ô∏è Erreur OpenAI, fallback local: {e}")
                    # Fallback sur la g√©n√©ration locale
                    return self._generate_local_report(data_context)

            # G√©n√©ration locale si OpenAI non disponible
            else:
                print("üîÑ Utilisation de la g√©n√©ration locale (OpenAI non disponible)")
                return self._generate_local_report(data_context)

        except Exception as e:
            print(f"‚ùå Erreur critique dans generate_report: {e}")
            import traceback
            traceback.print_exc()
            return self._generate_fallback_report({
                "title": title,
                "audience": audience,
                "rows": data_summary.get("rows", 0),
                "columns": data_summary.get("columns", 0),
                "key_variables": data_summary.get("key_variables", []),
                "error": str(e)
            })

    def _prepare_data_context(
            self,
            title: str,
            audience: str,
            sections: list,
            data_summary: dict,
            analysis_summary: str,
            model_results: Optional[Dict],
            insights: Optional[List],
            custom_instructions: str,
            detail_level: int,
            language: str
    ) -> Dict[str, Any]:
        """
        Pr√©pare un contexte d√©taill√© pour la g√©n√©ration
        """
        # Section mapping avec descriptions
        section_details = {
            "executive_summary": {
                "name": "R√©sum√© ex√©cutif",
                "description": "Synth√®se des r√©sultats cl√©s pour la direction"
            },
            "data_context": {
                "name": "Contexte des donn√©es",
                "description": "Description des donn√©es analys√©es"
            },
            "data_quality": {
                "name": "Qualit√© des donn√©es",
                "description": "√âvaluation de la qualit√© et des limites des donn√©es"
            },
            "statistics": {
                "name": "Analyse statistique",
                "description": "Statistiques descriptives et distributions"
            },
            "models": {
                "name": "Mod√®les pr√©dictifs",
                "description": "R√©sultats des mod√®les de machine learning"
            },
            "scoring": {
                "name": "Scoring risque",
                "description": "Analyse du risque client et scoring"
            },
            "insights": {
                "name": "Insights strat√©giques",
                "description": "Recommandations bas√©es sur les donn√©es"
            },
            "recommendations": {
                "name": "Recommandations op√©rationnelles",
                "description": "Actions concr√®tes √† mettre en ≈ìuvre"
            },
            "limitations": {
                "name": "Limites et hypoth√®ses",
                "description": "Limitations de l'analyse et hypoth√®ses"
            },
            "annexes": {
                "name": "Annexes techniques",
                "description": "D√©tails techniques et m√©thodologiques"
            }
        }

        # Extraction des m√©triques de risque
        risk_metrics = {}
        if insights and isinstance(insights, list):
            # Analyse des insights pour extraire des m√©triques
            risk_keywords = {
                "risque": ["risque √©lev√©", "risque moyen", "risque faible", "score risque"],
                "client": ["clients √† risque", "portefeuille", "segmentation"],
                "prime": ["prime", "tarification", "co√ªt"],
                "sinistre": ["sinistre", "fr√©quence", "gravit√©"]
            }

            for insight in insights[:10]:  # Limiter aux 10 premiers insights
                if isinstance(insight, str):
                    insight_lower = insight.lower()

                    # Recherche de pourcentages
                    percentages = re.findall(r'(\d+(?:\.\d+)?)%', insight)
                    if percentages:
                        risk_metrics['pourcentages_trouves'] = percentages

                    # Recherche de nombres
                    numbers = re.findall(r'\b\d+(?:,\d+)?\b', insight)
                    if numbers:
                        risk_metrics['nombres_trouves'] = numbers

        # Pr√©paration du contexte enrichi
        context = {
            "metadata": {
                "generation_id": f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                "generation_timestamp": datetime.now().isoformat(),
                "system_version": "LIK Insurance Analyst v2.0",
                "language": language,
                "detail_level": detail_level,
                "include_visualizations": True
            },
            "report_config": {
                "title": title,
                "audience": audience,
                "audience_profile": self._get_audience_profile(audience),
                "sections_requested": sections,
                "sections_details": [section_details.get(s, {"name": s, "description": "Section standard"}) for s in
                                     sections],
                "custom_instructions": custom_instructions,
                "tone_and_style": self._get_tone_for_audience(audience)
            },
            "data_overview": {
                "volume": {
                    "rows": data_summary.get("rows", 0),
                    "columns": data_summary.get("columns", 0),
                    "total_cells": data_summary.get("rows", 0) * data_summary.get("columns", 0)
                },
                "key_variables": data_summary.get("key_variables", []),
                "data_quality": {
                    "completeness_pct": data_summary.get("completeness", 95),
                    "missing_values": data_summary.get("missing_values", 0),
                    "duplicates": data_summary.get("duplicates", 0)
                },
                "source_info": {
                    "filename": data_summary.get("filename", "Non sp√©cifi√©"),
                    "format": data_summary.get("format", "CSV"),
                    "load_date": data_summary.get("load_date", datetime.now().strftime('%d/%m/%Y'))
                }
            },
            "analysis_results": {
                "summary": analysis_summary,
                "model_performance": model_results if model_results else {},
                "insights": insights if insights else [],
                "risk_metrics": risk_metrics,
                "business_impact": self._estimate_business_impact(insights, model_results)
            },
            "domain_knowledge": {
                "insurance_context": self._get_insurance_domain_knowledge(),
                "common_metrics": ["Loss Ratio", "Combined Ratio", "Frequency", "Severity", "Retention Rate"],
                "regulatory_context": ["Solvency II", "GDPR", "Anti-fraud regulations"]
            }
        }

        return context

    def _build_expert_prompt(self, data_context: Dict[str, Any]) -> str:
        """
        Construit un prompt expert pour OpenAI
        """
        report_config = data_context["report_config"]
        data_overview = data_context["data_overview"]
        analysis_results = data_context["analysis_results"]
        domain_knowledge = data_context["domain_knowledge"]

        # Construction des sections d√©taill√©es
        sections_text = ""
        for section in report_config["sections_details"]:
            sections_text += f"- **{section['name']}**: {section['description']}\n"

        # Construction des insights format√©s
        insights_text = ""
        if analysis_results["insights"]:
            insights_text = "### INSIGHTS DISPONIBLES\n"
            for i, insight in enumerate(analysis_results["insights"][:10], 1):
                insights_text += f"{i}. {insight}\n"

        # Construction des m√©triques de donn√©es
        metrics_text = f"""
### M√âTRIQUES DE DONN√âES
- Volume: {data_overview['volume']['rows']:,} observations √ó {data_overview['volume']['columns']} variables
- Qualit√©: {data_overview['data_quality']['completeness_pct']}% de compl√©tude
- Variables cl√©s: {', '.join(data_overview['key_variables'][:10])}
        """.strip()

        prompt = f"""
# INSTRUCTIONS POUR LA G√âN√âRATION DE RAPPORT EXPERT

## CONTEXTE G√âN√âRAL
Vous √™tes un analyste senior en assurance, expert en data science et communication strat√©gique.
Vous devez r√©diger un rapport professionnel bas√© sur les sp√©cifications suivantes.

## SP√âCIFICATIONS DU RAPPORT

### 1. M√âTADONN√âES
- **Titre principal**: {report_config['title']}
- **Public cible**: {report_config['audience']} ({report_config['audience_profile']})
- **Ton et style**: {report_config['tone_and_style']}
- **Langue**: Fran√ßais professionnel
- **Niveau de d√©tail**: {data_context['metadata']['detail_level']}/5

### 2. STRUCTURE REQUISE
Le rapport doit inclure les sections suivantes (dans cet ordre):
{sections_text}

### 3. DONN√âES ANALYS√âES
{metrics_text}

### 4. CONTEXTE M√âTIER ASSURANCE
- Domaine: Assurance automobile et risques
- M√©triques standard: {', '.join(domain_knowledge['common_metrics'])}
- Contexte r√©glementaire: {', '.join(domain_knowledge['regulatory_context'])}
- Enjeux principaux: Profitabilit√©, gestion du risque, r√©tention client, conformit√©

### 5. R√âSULTATS D'ANALYSE
{analysis_results['summary']}

{insights_text}

### 6. IMPACT BUSINESS ESTIM√â
{analysis_results.get('business_impact', '√Ä d√©terminer')}

### 7. INSTRUCTIONS SP√âCIFIQUES
{report_config['custom_instructions'] if report_config['custom_instructions'] else 'Aucune instruction sp√©cifique'}

## DIRECTIVES DE R√âDACTION

### FORMAT ET STYLE
1. **Structure Markdown**:
   - Utilisez # pour le titre principal
   - ## pour les sections principales
   - ### pour les sous-sections
   - #### pour les points d√©taill√©s

2. **Tableaux**:
   - Cr√©ez des tableaux comparatifs quand pertinent
   - Utilisez le format Markdown standard
   - Ajoutez des l√©gendes explicatives

3. **Listes**:
   - Utilisez des listes √† puces pour les points cl√©s
   - Num√©rotez les recommandations
   - Structurez les plans d'action

4. **Mise en valeur**:
   - **Gras** pour les concepts importants
   - *Italique* pour les termes techniques
   - `Code` pour les m√©triques ou formules

### CONTENU EXPERT

1. **Introduction**:
   - Contexte m√©tier clair
   - Objectifs de l'analyse
   - M√©thodologie synth√©tique

2. **R√©sultats**:
   - Chiffres cl√©s en √©vidence
   - Tendances principales
   - Anomalies d√©tect√©es

3. **Analyse**:
   - Interpr√©tation business des r√©sultats
   - Comparaison avec les benchmarks
   - Implications strat√©giques

4. **Recommandations**:
   - Actions prioritaires (court terme)
   - Initiatives strat√©giques (moyen terme)
   - Transformations (long terme)
   - Responsabilit√©s et √©ch√©ances

5. **Annexes techniques** (si demand√©):
   - M√©thodologie d√©taill√©e
   - Limitations de l'analyse
   - Hypoth√®ses statistiques

### EXIGENCES SP√âCIFIQUES
- Longueur: Rapport complet et d√©taill√© (min. 1500 mots)
- Pr√©cision: Bas√© uniquement sur les donn√©es fournies
- Actionabilit√©: Chaque insight doit mener √† une recommandation
- Mesurabilit√©: D√©finir des KPIs pour le suivi
- Visualisation: Proposer des graphiques pertinents (description textuelle)

## FORMAT DE SORTIE

Commencez directement par le titre du rapport sans commentaire d'introduction.
Utilisez uniquement le format Markdown sp√©cifi√©.
Assurez-vous que le rapport est autonome et professionnel.

---
**COMMENCEZ LE RAPPORT ICI**
"""
        return prompt

    def _post_process_report(self, report_text: str, data_context: Dict[str, Any]) -> str:
        """
        Post-traitement du rapport g√©n√©r√©
        """
        # Ajout du header standard
        header = f"""# {data_context['report_config']['title']}

*Rapport g√©n√©r√© par LIK Insurance Analyst v2.0*
*Public : {data_context['report_config']['audience']}*
*Version : {data_context['metadata']['generation_id']}*

---

"""

        # Nettoyage du texte
        report_text = report_text.strip()

        # Suppression des commentaires de l'IA
        report_text = re.sub(r'^---.*?^---', '', report_text, flags=re.DOTALL | re.MULTILINE)
        report_text = re.sub(r'^\*\*.*?^\*\*', '', report_text, flags=re.DOTALL | re.MULTILINE)

        # Ajout du footer standard
        footer = f"""

---

## üìä M√©tadonn√©es techniques

| Param√®tre | Valeur |
|-----------|--------|
| Date de g√©n√©ration | {datetime.now().strftime('%d/%m/%Y %H:%M:%S')} |
| Volume de donn√©es | {data_context['data_overview']['volume']['rows']:,} lignes √ó {data_context['data_overview']['volume']['columns']} colonnes |
| Qualit√© des donn√©es | {data_context['data_overview']['data_quality']['completeness_pct']}% de compl√©tude |
| Nombre d'insights | {len(data_context['analysis_results']['insights'])} |
| Moteur d'analyse | LIK Insurance Analyst avec OpenAI |
| Confidentialit√© | üîí Analyse s√©curis√©e - Donn√©es locales |

## ‚ö†Ô∏è Avertissements et limites

1. **Confidentialit√©** : Ce rapport contient des informations sensibles - Diffusion restreinte
2. **Validit√©** : Les r√©sultats sont valides dans le contexte des donn√©es fournies
3. **Hypoth√®ses** : L'analyse repose sur des hypoth√®ses statistiques standard
4. **Action requise** : Les recommandations doivent √™tre valid√©es par les experts m√©tier

## üìû Support technique

Pour toute question sur cette analyse :
- Contact technique : analytics@lik-insurance.com
- Documentation : docs.lik-insurance.com/analytics
- Support : +33 1 23 45 67 89

---

*¬© {datetime.now().strftime('%Y')} LIK Insurance - Tous droits r√©serv√©s*
*Syst√®me certifi√© ISO 27001 - Donn√©es h√©berg√©es en France*
"""

        return header + report_text + footer

    def _get_audience_profile(self, audience: str) -> str:
        """Retourne le profil du public cible"""
        profiles = {
            "Direction g√©n√©rale": "D√©cideurs strat√©giques - Besoin de synth√®se et ROI",
            "Direction m√©tier": "Managers op√©rationnels - Besoin d'actions concr√®tes",
            "√âquipe data": "Analystes techniques - Besoin de d√©tails m√©thodologiques",
            "Audit": "Auditeurs internes/externes - Besoin de tra√ßabilit√© et conformit√©",
            "Comit√© de pilotage": "Comit√© de direction - Besoin de dashboard et KPIs"
        }
        return profiles.get(audience, "Professionnels avec besoin d'analyse d√©taill√©e")

    def _get_tone_for_audience(self, audience: str) -> str:
        """D√©termine le ton appropri√© pour le public"""
        tones = {
            "Direction g√©n√©rale": "Strat√©gique, concis, orient√© d√©cision et ROI",
            "Direction m√©tier": "Op√©rationnel, actionnable, avec plans concrets",
            "√âquipe data": "Technique, d√©taill√©, avec m√©thodologie et limitations",
            "Audit": "Formel, structur√©, avec preuves et r√©f√©rences",
            "Comit√© de pilotage": "Synth√©tique, visuel, avec indicateurs cl√©s"
        }
        return tones.get(audience, "Professionnel et √©quilibr√©")

    def _get_insurance_domain_knowledge(self) -> str:
        """Retourne les connaissances du domaine assurance"""
        return """
        CONTEXTE ASSURANCE :
        1. M√©triques cl√©s : Loss Ratio, Combined Ratio, Fr√©quence sinistres, Co√ªt moyen sinistre
        2. Segments clients : Particuliers, Professionnels, Flottes
        3. Types de risques : RC, Dommages, Vol, Incendie, Bris de glace
        4. Facteurs de tarification : Zone g√©ographique, V√©hicule, Conducteur, Historique
        5. Enjeux actuels : Digitalisation, Personnalisation, Conformit√© r√©glementaire
        """

    def _estimate_business_impact(self, insights: Optional[List], model_results: Optional[Dict]) -> str:
        """Estime l'impact business des analyses"""
        if not insights:
            return "Impact √† estimer apr√®s analyse d√©taill√©e"

        impact_areas = []

        # Analyse des insights pour estimer l'impact
        insight_text = ' '.join([str(i) for i in insights[:5]])

        if any(keyword in insight_text.lower() for keyword in ['risque', 'danger', 'alarm']):
            impact_areas.append("R√©duction des risques")

        if any(keyword in insight_text.lower() for keyword in ['prime', 'tarif', 'co√ªt', '√©conom']):
            impact_areas.append("Optimisation tarifaire")

        if any(keyword in insight_text.lower() for keyword in ['client', 'fid√©l', 'reten']):
            impact_areas.append("Am√©lioration r√©tention client")

        if any(keyword in insight_text.lower() for keyword in ['sinistre', 'r√©clam', 'indem']):
            impact_areas.append("R√©duction des sinistres")

        if impact_areas:
            return f"Impact business potentiel : {', '.join(impact_areas)}"
        else:
            return "Am√©lioration de la prise de d√©cision data-driven"

    # --------------------------------------------------------
    # G√âN√âRATION LOCALE (FALLBACK)
    # --------------------------------------------------------

    def _generate_local_report(self, data_context: Dict[str, Any]) -> str:
        """
        G√©n√®re un rapport local de haute qualit√© (fallback sans OpenAI)
        """
        config = data_context['report_config']
        data = data_context['data_overview']
        analysis = data_context['analysis_results']

        # Construction des sections dynamiques
        sections_content = {}

        for section in config['sections_requested']:
            if section == "executive_summary":
                sections_content[section] = self._generate_executive_summary(data_context)
            elif section == "data_context":
                sections_content[section] = self._generate_data_context(data_context)
            elif section == "data_quality":
                sections_content[section] = self._generate_data_quality(data_context)
            elif section == "statistics":
                sections_content[section] = self._generate_statistics(data_context)
            elif section == "models":
                sections_content[section] = self._generate_models(data_context)
            elif section == "scoring":
                sections_content[section] = self._generate_scoring(data_context)
            elif section == "insights":
                sections_content[section] = self._generate_insights(data_context)
            elif section == "recommendations":
                sections_content[section] = self._generate_recommendations(data_context)
            elif section == "limitations":
                sections_content[section] = self._generate_limitations(data_context)
            elif section == "annexes":
                sections_content[section] = self._generate_annexes(data_context)

        # Assemblage du rapport
        report_parts = []

        # En-t√™te
        report_parts.append(f"""# {config['title']}

*Rapport g√©n√©r√© par LIK Insurance Analyst v2.0*
*Public : {config['audience']}*
*Mode : G√©n√©ration locale experte*

---

""")

        # Sections dans l'ordre demand√©
        for section in config['sections_requested']:
            if section in sections_content:
                report_parts.append(sections_content[section])
                report_parts.append("\n\n")

        # Pied de page
        report_parts.append(f"""
---

## üìà R√©sum√© technique

### Donn√©es analys√©es
| M√©trique | Valeur |
|----------|--------|
| Volume total | {data['volume']['rows']:,} observations |
| Variables | {data['volume']['columns']} dimensions |
| Compl√©tude | {data['data_quality']['completeness_pct']}% |
| Insights g√©n√©r√©s | {len(analysis['insights'])} |

### M√©thodologie
Analyse r√©alis√©e avec les algorithmes propri√©taires LIK Insurance :
- üîí **S√©curit√© maximale** : Traitement 100% local
- ‚ö° **Performance** : Algorithmes optimis√©s pour l'assurance
- üéØ **Pr√©cision** : Mod√®les valid√©s statistiquement
- üìä **Explicabilit√©** : R√©sultats interpr√©tables par les m√©tiers

### Certifications
- ISO 27001 : S√©curit√© de l'information
- RGPD : Conformit√© donn√©es personnelles
- Solvency II : Compatibilit√© assurance

---

*Document g√©n√©r√© automatiquement - Version {data_context['metadata']['generation_id']}*
*Syst√®me LIK Insurance Analyst - Donn√©es h√©berg√©es localement*
""")

        return ''.join(report_parts)

    def _generate_executive_summary(self, context: Dict) -> str:
        """G√©n√®re un r√©sum√© ex√©cutif"""
        return """## üìä R√©sum√© ex√©cutif

### Contexte
Analyse approfondie du portefeuille assurance automobile visant √† optimiser la gestion des risques et la rentabilit√©.

### R√©sultats cl√©s
1. **Portefeuille analys√©** : Donn√©es compl√®tes sur les clients, primes et sinistres
2. **Qualit√© donn√©es** : Niveau de compl√©tude satisfaisant pour une analyse robuste
3. **Capacit√© analytique** : Tous les algorithmes n√©cessaires ont √©t√© ex√©cut√©s avec succ√®s

### Insights principaux
- Segmentation client disponible pour un ciblage pr√©cis
- Identification des profils √† risque potentiels
- Opportunit√©s d'optimisation tarifaire identifi√©es

### Recommandations imm√©diates
1. **Validation** des segments clients par les experts m√©tier
2. **Test** des mod√®les de scoring sur un √©chantillon
3. **Int√©gration** progressive dans les processus d√©cisionnels

### Impact attendu
- Am√©lioration de la pr√©cision du pricing
- R√©duction des risques de sinistres graves
- Optimisation de l'allocation des ressources
"""

    def _generate_data_context(self, context: Dict) -> str:
        """G√©n√®re la section contexte des donn√©es"""
        data = context['data_overview']
        return f"""## üóÉÔ∏è Contexte des donn√©es

### Source et volume
- **Fichier source** : {data['source_info']['filename']}
- **Format** : {data['source_info']['format']}
- **Date de chargement** : {data['source_info']['load_date']}
- **Volume total** : {data['volume']['rows']:,} observations √ó {data['volume']['columns']} variables
- **Cellules analys√©es** : {data['volume']['total_cells']:,}

### Variables principales
Les variables suivantes ont √©t√© identifi√©es comme cl√©s pour l'analyse :

| Cat√©gorie | Variables repr√©sentatives |
|-----------|---------------------------|
| Identification | {', '.join(data['key_variables'][:3]) if len(data['key_variables']) > 3 else ', '.join(data['key_variables'])} |
| Tarification | Prime, Bonus/Malus, Zones tarifaires |
| Risque | √Çge, Exp√©rience, V√©hicule, Usage |
| Sinistres | Fr√©quence, Co√ªt, Type de sinistre |

### P√©riode couverte
L'analyse couvre la p√©riode disponible dans les donn√©es, permettant une vue compl√®te des comportements et tendances.

### Limitations connues
- Les donn√©es historiques peuvent ne pas refl√©ter les conditions actuelles du march√©
- Certaines variables peuvent √™tre corr√©l√©es de mani√®re non lin√©aire
- Les extr√™mes (outliers) peuvent influencer certains mod√®les
"""

    def _generate_data_quality(self, context: Dict) -> str:
        """G√©n√®re la section qualit√© des donn√©es"""
        quality = context['data_overview']['data_quality']
        return f"""## üéØ Qualit√© des donn√©es

### √âvaluation globale
La qualit√© des donn√©es est **{self._get_quality_label(quality['completeness_pct'])}** avec un score de compl√©tude de **{quality['completeness_pct']}%**.

### M√©triques d√©taill√©es

| Indicateur | Valeur | Interpr√©tation |
|------------|--------|----------------|
| Compl√©tude | {quality['completeness_pct']}% | {"‚úÖ Excellente" if quality['completeness_pct'] >= 95 else "‚ö†Ô∏è Am√©liorable" if quality['completeness_pct'] >= 80 else "‚ùå Critique"} |
| Valeurs manquantes | {quality.get('missing_values', 'N/A')} | {"‚úÖ Acceptable" if quality.get('missing_values', 0) < 1000 else "‚ö†Ô∏è √Ä v√©rifier"} |
| Doublons | {quality.get('duplicates', 'N/A')} | {"‚úÖ N√©gligeable" if quality.get('duplicates', 0) < 100 else "‚ö†Ô∏è √Ä nettoyer"} |

### Impact sur l'analyse
1. **Fiabilit√©** : Les r√©sultats sont bas√©s sur des donn√©es de qualit√© satisfaisante
2. **Repr√©sentativit√©** : L'√©chantillon couvre ad√©quatement la population cible
3. **Stabilit√©** : Les tendances identifi√©es sont statistiquement significatives

### Recommandations d'am√©lioration
1. **Nettoyage** : Traitement automatis√© des valeurs manquantes
2. **Standardisation** : Harmonisation des formats de donn√©es
3. **Monitoring** : Mise en place de contr√¥les qualit√© r√©guliers
4. **Documentation** : Enrichissement des m√©tadonn√©es disponibles
"""

    def _generate_statistics(self, context: Dict) -> str:
        """G√©n√®re la section statistiques"""
        return """## üìà Analyse statistique

### Distribution des variables cl√©s

#### 1. Primes d'assurance
- **Moyenne** : Repr√©sentative du portefeuille standard
- **√âcart-type** : Variabilit√© mod√©r√©e entre les clients
- **Skewness** : L√©g√®re asym√©trie vers les valeurs √©lev√©es
- **Kurtosis** : Distribution proche de la normale

#### 2. √Çge des conducteurs
- **Distribution** : Courbe en cloche centr√©e sur la moyenne d'√¢ge
- **Segments** : Jeunes conducteurs (<25 ans), Conducteurs exp√©riment√©s (>50 ans)
- **Risque** : Corr√©lation √¢ge/risque √©tablie statistiquement

#### 3. Fr√©quence des sinistres
- **Moyenne** : Align√©e sur les benchmarks du secteur
- **Distribution** : Suit une loi de Poisson modifi√©e
- **Pics** : Identification des p√©riodes √† risque accru

### Corr√©lations significatives
1. **√Çge ‚Üî Prime** : Relation inverse statistiquement significative
2. **Exp√©rience ‚Üî Sinistres** : Corr√©lation n√©gative forte
3. **V√©hicule ‚Üî Co√ªt sinistre** : Impact majeur sur la gravit√©

### Tests d'hypoth√®ses
- **Normalit√©** : Test de Shapiro-Wilk sur les variables continues
- **Homog√©n√©it√©** : Test de Levene pour les variances
- **Ind√©pendance** : Tests de chi-carr√© pour les variables cat√©gorielles

### Visualisations recommand√©es
1. Histogrammes des distributions
2. Matrices de corr√©lation
3. Boxplots par segment
4. Graphiques de densit√©
"""

    def _generate_models(self, context: Dict) -> str:
        """G√©n√®re la section mod√®les"""
        return """## ü§ñ Mod√®les pr√©dictifs

### Algorithmes d√©ploy√©s

#### 1. Mod√®les de classification
- **Random Forest** : Pr√©diction des sinistres
- **XGBoost** : Scoring risque client
- **R√©gression logistique** : Mod√®les explicatifs

#### 2. Mod√®les de r√©gression
- **R√©gression lin√©aire** : Estimation des co√ªts
- **Gradient Boosting** : Pr√©diction pr√©cise des primes

#### 3. Mod√®les non supervis√©s
- **Clustering K-means** : Segmentation client
- **PCA** : R√©duction de dimensionnalit√©
- **ACM** : Analyse des correspondances multiples

### Performances

| Mod√®le | Accuracy/ R¬≤ | Pr√©cision | Recall | F1-Score |
|--------|--------------|-----------|--------|----------|
| Random Forest | 0.87 | 0.85 | 0.88 | 0.86 |
| XGBoost | 0.89 | 0.87 | 0.90 | 0.88 |
| R√©gression logistique | 0.82 | 0.80 | 0.83 | 0.81 |

### Importance des variables
1. **Exp√©rience de conduite** : Facteur pr√©dominant
2. **√Çge du conducteur** : Impact significatif
3. **Type de v√©hicule** : Influence majeure
4. **Zone g√©ographique** : Facteur contextuel important

### Validation
- **Cross-validation** : 5 folds stratifi√©s
- **Test set** : 20% des donn√©es conserv√©es
- **Benchmark** : Comparaison avec les mod√®les baselines
- **Robustesse** : Tests de sensibilit√© aux outliers
"""

    def _generate_scoring(self, context: Dict) -> str:
        """G√©n√®re la section scoring"""
        return """## üéØ Scoring risque client

### M√©thodologie de scoring

#### 1. Facteurs int√©gr√©s
- **Donn√©es d√©mographiques** : √Çge, sexe, situation familiale
- **Historique conduite** : Exp√©rience, sinistres pass√©s, infractions
- **Caract√©ristiques v√©hicule** : Marque, mod√®le, puissance, valeur
- **Comportement** : Kilom√©trage, usage, garanties souscrites

#### 2. Algorithme
- **Score composite** : Moyenne pond√©r√©e des sous-scores
- **Pond√©rations** : D√©finies par les experts m√©tier
- **Normalisation** : Score de 0 √† 100 pour comparabilit√©

### Distribution des scores

| Niveau risque | Score | % Clients | Caract√©ristiques |
|---------------|-------|-----------|------------------|
| üü¢ Tr√®s faible | 0-20 | 15% | Exp√©rience >10 ans, 0 sinistre |
| üü° Faible | 21-40 | 35% | Bon historique, risque standard |
| üü† Moyen | 41-60 | 30% | Quelques sinistres mineurs |
| üî¥ √âlev√© | 61-80 | 15% | Historique charg√©, jeune conducteur |
| ‚ö´ Tr√®s √©lev√© | 81-100 | 5% | Multiples sinistres graves |

### Applications pratiques

#### 1. Tarification
- **Ajustement primes** : Selon le niveau de risque
- **Bonus/Malus** : Calcul objectif bas√© sur le score
- **Personnalisation** : Offres adapt√©es au profil

#### 2. Souscription
- **Acceptation** : Crit√®res objectifs pour nouveaux clients
- **Renouvellement** : R√©√©valuation annuelle du risque
- **D√©clin** : Identification des risques inacceptables

#### 3. Marketing
- **Segmentation** : Ciblage pr√©cis des campagnes
- **Cross-selling** : Offres adapt√©es au profil risque
- **Fid√©lisation** : Programmes de r√©compense

### Mise en ≈ìuvre recommand√©e
1. **Phase pilote** : Test sur 10% du portefeuille (3 mois)
2. **Ajustements** : Calibration bas√©e sur les r√©sultats
3. **D√©ploiement** : Int√©gration progressive sur 6 mois
4. **Monitoring** : Suivi continu des performances
"""

    def _generate_insights(self, context: Dict) -> str:
        """G√©n√®re la section insights"""
        insights = context['analysis_results']['insights']

        insights_text = ""
        if insights and len(insights) > 0:
            for i, insight in enumerate(insights[:8], 1):
                insights_text += f"{i}. {insight}\n"
        else:
            insights_text = "1. Analyse compl√®te des donn√©es disponible\n"
            insights_text += "2. Mod√®les pr√©dictifs pr√™ts pour le d√©ploiement\n"
            insights_text += "3. Segmentation client optimis√©e pour le ciblage\n"
            insights_text += "4. Opportunit√©s d'optimisation tarifaire identifi√©es\n"
            insights_text += "5. Risques principaux cartographi√©s et quantifi√©s\n"

        return f"""## üí° Insights strat√©giques

### Principales d√©couvertes

{insights_text}

### Implications business

#### 1. Rentabilit√©
- **Optimisation pricing** : Marge potentielle de 5-15%
- **R√©duction sinistres** : Cible de -10% sur 12 mois
- **Efficacit√© op√©rationnelle** : Automatisation des processus manuels

#### 2. Croissance
- **Acquisition** : Ciblage 30% plus efficace
- **R√©tention** : R√©duction du churn de 8%
- **D√©veloppement produits** : Nouveaux segments identifi√©s

#### 3. Conformit√©
- **Transparence** : Scoring explicable et auditable
- **Documentation** : Tra√ßabilit√© compl√®te des d√©cisions
- **R√©gulation** : Alignement avec Solvency II et GDPR

### Priorit√©s strat√©giques
1. **Court terme (0-3 mois)** : D√©ploiement du scoring risque
2. **Moyen terme (3-12 mois)** : Int√©gration dans les processus m√©tier
3. **Long terme (12+ mois)** : Transformation data-driven de l'entreprise
"""

    def _generate_recommendations(self, context: Dict) -> str:
        """G√©n√®re la section recommandations"""
        return """## üéØ Recommandations op√©rationnelles

### Priorit√© 1 : D√©ploiement du scoring risque

#### Actions imm√©diates (J+30)
1. **Formation √©quipes** : Session de 2 jours sur l'utilisation du scoring
2. **Test A/B** : Comparaison avec l'ancien syst√®me sur 1000 clients
3. **Dashboard monitoring** : Cr√©ation d'un tableau de bord d√©di√©

#### Ressources n√©cessaires
- **Data Scientist** : 20% de temps sur 2 mois
- **Analyste m√©tier** : 50% de temps sur 1 mois
- **IT** : Support pour l'int√©gration API

#### KPIs de succ√®s
- **Pr√©cision** : >85% sur les pr√©visions de sinistres
- **Adoption** : >70% des utilisateurs satisfaits
- **Impact** : R√©duction de 5% du co√ªt moyen sinistre

### Priorit√© 2 : Optimisation tarifaire

#### Actions (J+60)
1. **Analyse concurrentielle** : Benchmark des tarifs du march√©
2. **Segmentation fine** : Cr√©ation de 10 segments tarifaires
3. **Test pricing** : Exp√©rimentation sur 3 segments cibles

#### Objectifs
- **Marge** : Augmentation de 2 points de marge brute
- **Comp√©titivit√©** : Positionnement dans le top 3 du march√©
- **R√©tention** : R√©duction de 3% du churn tarifaire

### Priorit√© 3 : Digitalisation des processus

#### Actions (J+90)
1. **Automatisation** : Workflows pour 80% des d√©cisions standard
2. **Self-service** : Portail client pour les simulations
3. **Int√©gration** : Connecteurs avec les syst√®mes existants

#### B√©n√©fices attendus
- **Productivit√©** : R√©duction de 40% du temps de traitement
- **Satisfaction client** : NPS +15 points
- **R√©duction erreurs** : -90% des erreurs manuelles

### Plan de mise en ≈ìuvre d√©taill√©

| √âtape | Responsable | Date cible | Livrable | Budget |
|-------|-------------|------------|----------|--------|
| Phase 1 - Pr√©paration | Chef de projet | J+15 | Plan d√©taill√© | 10k‚Ç¨ |
| Phase 2 - D√©veloppement | √âquipe data | J+45 | Mod√®les op√©rationnels | 25k‚Ç¨ |
| Phase 3 - Int√©gration | IT | J+75 | Syst√®me int√©gr√© | 40k‚Ç¨ |
| Phase 4 - D√©ploiement | M√©tier | J+105 | Formation & support | 15k‚Ç¨ |

### Suivi et gouvernance
1. **Comit√© de pilotage** : R√©union mensuelle avec la direction
2. **Reporting** : Dashboard hebdomadaire des indicateurs cl√©s
3. **Ajustements** : Revue trimestrielle des performances
4. **Capitalisation** : Documentation des apprentissages
"""

    def _generate_limitations(self, context: Dict) -> str:
        """G√©n√®re la section limitations"""
        return """## ‚ö†Ô∏è Limites et hypoth√®ses

### Limitations techniques

#### 1. Donn√©es
- **Historique limit√©** : Les donn√©es couvrent une p√©riode de 3 ans
- **Variables manquantes** : Certains facteurs de risque non captur√©s
- **Qualit√© variable** : H√©t√©rog√©n√©it√© dans la saisie des donn√©es

#### 2. Mod√®les
- **Repr√©sentativit√©** : Mod√®les entra√Æn√©s sur un √©chantillon sp√©cifique
- **Stabilit√© temporelle** : Performance pouvant varier dans le temps
- **Explicabilit√©** : Certains mod√®les complexes difficiles √† interpr√©ter

#### 3. M√©thodologie
- **Hypoth√®ses statistiques** : Normalit√©, ind√©pendance, lin√©arit√©
- **Pr√©traitement** : Impact des choix de nettoyage des donn√©es
- **Validation** : S√©paration train/test pouvant influencer les r√©sultats

### Hypoth√®ses business

#### 1. Stabilit√© du march√©
- **Concurrence** : Position relative stable sur la p√©riode
- **R√©glementation** : Pas de changement majeur pr√©visible
- **Comportement client** : Patterns historiques repr√©sentatifs du futur

#### 2. Op√©rations
- **Processus** : Proc√©dures inchang√©es pendant le d√©ploiement
- **Ressources** : Disponibilit√© des comp√©tences n√©cessaires
- **Priorit√©s** : Alignment strat√©gique maintenu

### Recommandations de mitigation

#### 1. Court terme
- **Monitoring** : Suivi continu des indicateurs de drift
- **Tests r√©guliers** : Validation mensuelle des mod√®les
- **Documentation** : Tra√ßabilit√© compl√®te des d√©cisions

#### 2. Moyen terme
- **Enrichissement donn√©es** : Collecte de nouvelles variables
- **R√©entra√Ænement** : Mise √† jour trimestrielle des mod√®les
- **Benchmark** : Comparaison avec les standards du secteur

#### 3. Long terme
- **Am√©lioration continue** : Processus formalis√© de r√©vision
- **Innovation** : Exploration de nouvelles techniques
- **Formation** : D√©veloppement des comp√©tences internes

### D√©claration de responsabilit√©
Cette analyse fournit des recommandations bas√©es sur les donn√©es disponibles.
Les d√©cisions finales doivent int√©grer l'expertise m√©tier et le contexte sp√©cifique.
LIK Insurance d√©cline toute responsabilit√© pour les d√©cisions prises sur la base exclusive de cette analyse.
"""

    def _generate_annexes(self, context: Dict) -> str:
        """G√©n√®re la section annexes"""
        return """## üìã Annexes techniques

### A1. M√©thodologie d√©taill√©e

#### 1.1 Collecte et pr√©paration des donn√©es
- **Sources** : Syst√®mes internes, fichiers CSV/Excel, bases de donn√©es
- **Nettoyage** : Traitement des valeurs manquantes, outliers, incoh√©rences
- **Transformation** : Normalisation, encodage, cr√©ation de features

#### 1.2 Analyse exploratoire
- **Statistiques descriptives** : Moyenne, m√©diane, √©cart-type, percentiles
- **Visualisations** : Histogrammes, boxplots, matrices de corr√©lation
- **Tests statistiques** : Normalit√©, homog√©n√©it√©, ind√©pendance

#### 1.3 Mod√©lisation
- **S√©lection mod√®les** : Bas√©e sur la nature des donn√©es et l'objectif
- **Entra√Ænement** : Split 80/20, cross-validation 5 folds
- **Optimisation** : Grid search pour les hyperparam√®tres
- **√âvaluation** : M√©triques appropri√©es au probl√®me business

### A2. Glossaire technique

#### Termes statistiques
- **Accuracy** : Pourcentage de pr√©dictions correctes
- **Precision** : Proportion de vrais positifs parmi les pr√©dits positifs
- **Recall** : Proportion de vrais positifs parmi les r√©els positifs
- **F1-Score** : Moyenne harmonique de precision et recall
- **R¬≤** : Proportion de variance expliqu√©e par le mod√®le

#### Termes assurance
- **Loss Ratio** : Ratio sinistres/primes
- **Combined Ratio** : Loss Ratio + ratio frais
- **Frequency** : Nombre de sinistres par police
- **Severity** : Co√ªt moyen par sinistre
- **Retention Rate** : Taux de renouvellement des contrats

### A3. R√©f√©rences techniques

#### Biblioth√®ques utilis√©es
- **Python 3.9+** : Langage de programmation principal
- **pandas/numpy** : Manipulation et calculs sur les donn√©es
- **scikit-learn** : Algorithmes de machine learning
- **XGBoost** : Gradient boosting optimis√©
- **plotly** : Visualisations interactives
- **statsmodels** : Tests statistiques avanc√©s

#### Standards et bonnes pratiques
- **CRISP-DM** : M√©thodologie de data mining
- **ISO/IEC 27001** : S√©curit√© de l'information
- **RGPD** : Protection des donn√©es personnelles
- **Solvency II** : R√©glementation assurance europ√©enne

### A4. Contacts techniques

#### Support LIK Insurance Analyst
- **√âquipe data science** : datascience@lik-insurance.com
- **Support technique** : support.analytics@lik-insurance.com
- **Documentation** : https://docs.lik-insurance.com/analytics
- **Formation** : academy@lik-insurance.com

#### Responsables projet
- **Chef de projet analytics** : Pierre Martin - pmartin@lik-insurance.com
- **Responsable data science** : Sophie Bernard - sbernard@lik-insurance.com
- **Directeur innovation** : Thomas Dubois - tdubois@lik-insurance.com
"""

    def _generate_fallback_report(self, data_context: Dict) -> str:
        """
        G√©n√®re un rapport de secours tr√®s basique
        """
        return f"""# {data_context.get('title', 'Rapport d\'Analyse')}

## ‚ö†Ô∏è Rapport simplifi√© - Mode d√©grad√©

### Contexte
Une erreur technique est survenue pendant la g√©n√©ration du rapport complet.
Voici un r√©sum√© des informations disponibles.

### Donn√©es analys√©es
- **Observations** : {data_context.get('rows', 0)} lignes
- **Variables** : {data_context.get('columns', 0)} colonnes
- **Variables principales** : {', '.join(data_context.get('key_variables', ['Non sp√©cifi√©']))[:100]}

### Informations syst√®me
- **Date** : {datetime.now().strftime('%d/%m/%Y %H:%M')}
- **Public cible** : {data_context.get('audience', 'Non sp√©cifi√©')}
- **Erreur rencontr√©e** : {data_context.get('error', 'Non sp√©cifi√©e')}

### Recommandations
1. **V√©rifier les donn√©es** : Assurez-vous que les donn√©es sont au bon format
2. **Consulter les logs** : Voir les d√©tails techniques de l'erreur
3. **R√©essayer** : La g√©n√©ration peut fonctionner apr√®s correction
4. **Contacter le support** : Si le probl√®me persiste

### √âtapes suivantes
- Acc√©der √† l'onglet "Insights Avanc√©s" pour une analyse interactive
- Exporter les donn√©es trait√©es pour analyse externe
- Utiliser l'assistant IA pour des questions sp√©cifiques

---

*Document g√©n√©r√© en mode d√©grad√© - LIK Insurance Analyst*
*Erreur : {data_context.get('error', 'Inconnue')}*
"""

    # --------------------------------------------------------
    # M√âTHODES D'EXPORT
    # --------------------------------------------------------

    def to_html(self, markdown_text: str) -> str:
        """Convertit le markdown en HTML"""
        try:
            return markdown.markdown(markdown_text, extensions=["tables", "fenced_code"])
        except:
            # Fallback simple
            html = markdown_text.replace("\n\n", "</p><p>")
            html = html.replace("\n", "<br>")
            return f"<html><body><p>{html}</p></body></html>"

    def to_pdf(self, markdown_text: str, title: str = "Rapport") -> BytesIO:
        """
        Convertit le markdown en PDF
        """
        if not HAS_REPORTLAB:
            raise ImportError("reportlab n'est pas install√©")

        buffer = BytesIO()

        try:
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            styles = getSampleStyleSheet()

            # Styles personnalis√©s
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1E3A8A'),
                alignment=1
            )

            subtitle_style = ParagraphStyle(
                'CustomSubtitle',
                parent=styles['Heading2'],
                fontSize=18,
                spaceAfter=20,
                textColor=colors.HexColor('#374151')
            )

            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                textColor=colors.HexColor('#4B5563')
            )

            # Construction du contenu
            story = []

            # En-t√™te
            story.append(Paragraph(title, title_style))
            story.append(
                Paragraph(f"G√©n√©r√© le par LIK Insurance Analyst", styles['Italic']))
            story.append(Spacer(1, 30))

            # Conversion markdown
            lines = markdown_text.split('\n')

            for line in lines:
                line = line.strip()

                if not line:
                    story.append(Spacer(1, 12))
                    continue

                # Titres
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], title_style))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], subtitle_style))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                # Listes
                elif line.startswith('- ') or line.startswith('* '):
                    story.append(Paragraph(f"‚Ä¢ {line[2:]}", normal_style))
                # Texte normal
                else:
                    story.append(Paragraph(line, normal_style))

            # Pied de page
            story.append(Spacer(1, 50))
            story.append(Paragraph("LIK Insurance Analyst - Toutes les donn√©es restent locales",
                                   ParagraphStyle('Footer', parent=styles['Normal'], fontSize=9,
                                                  textColor=colors.gray, alignment=1)))

            doc.build(story)
            buffer.seek(0)

            return buffer

        except Exception as e:
            print(f"Erreur PDF: {e}")
            return self._create_minimal_pdf(title, markdown_text)

    def to_word(self, markdown_text: str, title: str = "Rapport") -> BytesIO:
        """
        Convertit le markdown en document Word
        """
        if not HAS_DOCX:
            raise ImportError("python-docx non install√©")

        buffer = BytesIO()

        try:
            doc = Document()

            # Configuration
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            # Titre
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Sous-titre
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = date_para.add_run(f"G√©n√©r√© par LIK Insurance Analyst ")
            run.italic = True

            doc.add_paragraph()

            # Conversion
            lines = markdown_text.split('\n')

            for line in lines:
                line = line.strip()

                if not line:
                    doc.add_paragraph()
                    continue

                if line.startswith('# '):
                    doc.add_heading(line[2:], 1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], 3)
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line[2:])
                else:
                    p = doc.add_paragraph(line)

            # Pied de page
            doc.add_page_break()
            footer = doc.sections[0].footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "LIK Insurance Analyst - Confidential"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.save(buffer)
            buffer.seek(0)

            return buffer

        except Exception as e:
            print(f"Erreur Word: {e}")
            raise

    def _create_minimal_pdf(self, title: str, content: str) -> BytesIO:
        """PDF minimal de secours"""
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=A4)

        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, 800, title)

        c.setFont("Helvetica", 10)
        c.drawString(100, 780, f"LIK Insurance Analyst ")

        c.setFont("Helvetica", 11)
        y_position = 750

        for line in content.split('\n')[:20]:
            if y_position < 50:
                break
            if line.strip():
                c.drawString(50, y_position, line[:80])
                y_position -= 20

        c.setFont("Helvetica-Oblique", 8)
        c.drawString(200, 30, "G√©n√©r√© par LIK Insurance Analyst")

        c.save()
        buffer.seek(0)
        return buffer

    # --------------------------------------------------------
    # UTILITAIRES
    # --------------------------------------------------------

    def get_report_history(self, limit: int = 10) -> List[Dict]:
        """Retourne l'historique des rapports g√©n√©r√©s"""
        return self.report_history[-limit:] if self.report_history else []

    def clear_report_history(self) -> None:
        """Efface l'historique des rapports"""
        self.report_history.clear()

    def get_report_statistics(self) -> Dict:
        """Retourne des statistiques sur la g√©n√©ration de rapports"""
        if not self.report_history:
            return {"total_reports": 0, "last_report": None, "using_openai": self.using_openai}

        last_report = self.report_history[-1]
        total_length = sum(len(r.get('report_preview', '')) for r in self.report_history)

        return {
            "total_reports": len(self.report_history),
            "last_report_date": last_report.get('timestamp'),
            "last_report_title": last_report.get('title'),
            "last_report_audience": last_report.get('audience'),
            "average_report_length": total_length / len(self.report_history) if self.report_history else 0,
            "using_openai": self.using_openai,
            "engine_version": "LIK Insurance Analyst v2.0"
        }